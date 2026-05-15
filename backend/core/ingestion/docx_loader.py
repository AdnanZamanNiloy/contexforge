from __future__ import annotations

import asyncio
import logging
from typing import List, Optional

from .base_loader import BaseLoader
from core.types import Document
from observability.tracer import observe

logger = logging.getLogger(__name__)


class DocxLoader(BaseLoader):

    @observe(name="load_docx")
    async def load(
        self,
        source: str | bytes,
        source_id: str,
        filename: Optional[str] = None,
    ) -> List[Document]:
        text = await asyncio.to_thread(self._extract_text, source)
        return [
            Document(
                document_id=source_id,
                text=text,
                metadata={
                    "source_id": source_id,
                    "source_type": "docx",
                    "filename": filename or "unknown.docx",
                },
                source_type="docx",
            )
        ]

    def _extract_text(self, source: str | bytes) -> str:
        try:
            import docx
        except ImportError as exc: 
            logger.exception("python-docx not installed")
            raise RuntimeError("DOCX loader requires python-docx") from exc

        if isinstance(source, bytes):
            from io import BytesIO
            document = docx.Document(BytesIO(source))
        else:
            document = docx.Document(source)

        return "\n".join(paragraph.text for paragraph in document.paragraphs)